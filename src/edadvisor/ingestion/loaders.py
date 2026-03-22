from __future__ import annotations

import re
from pathlib import Path

import httpx
from bs4 import BeautifulSoup
from langchain_core.documents import Document
from loguru import logger


def _clean(text: str) -> str:
    """Strip excessive whitespace while preserving paragraph breaks."""
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def load_pdf(path: Path) -> list[Document]:
    """
    Extract text from a PDF using PyMuPDF.
    Each page becomes a separate Document with page number in metadata.
    """
    try:
        import fitz  # pymupdf
    except ImportError:
        raise ImportError("pip install pymupdf")

    docs = []
    with fitz.open(str(path)) as pdf:
        for i, page in enumerate(pdf):
            text = page.get_text("text")
            if not text.strip():
                continue
            docs.append(Document(
                page_content=_clean(text),
                metadata={
                    "source": path.name,
                    "source_path": str(path),
                    "page": i + 1,
                    "doc_type": "pdf",
                },
            ))
    logger.info(f"PDF  {path.name}  →  {len(docs)} pages")
    return docs


def load_html(url_or_path: str, doc_type: str = "web") -> list[Document]:
    """
    Load an HTML page — either a local file or a URL.
    Strips nav/header/footer boilerplate and returns clean body text.
    """
    if url_or_path.startswith("http"):
        try:
            resp = httpx.get(url_or_path, timeout=20, follow_redirects=True)
            resp.raise_for_status()
            raw_html = resp.text
            source = url_or_path
        except Exception as e:
            logger.error(f"fetch failed {url_or_path}: {e}")
            return []
    else:
        p = Path(url_or_path)
        raw_html = p.read_text(encoding="utf-8", errors="replace")
        source = p.name

    soup = BeautifulSoup(raw_html, "html.parser")

    # kill nav, scripts, styles, ads
    for tag in soup.find_all(["nav", "footer", "header", "script", "style", "aside"]):
        tag.decompose()

    body = soup.find("main") or soup.find("article") or soup.find("body")
    text = body.get_text(separator="\n") if body else soup.get_text(separator="\n")

    cleaned = _clean(text)
    if not cleaned:
        return []

    doc = Document(
        page_content=cleaned,
        metadata={"source": source, "doc_type": doc_type},
    )
    logger.info(f"HTML  {source}  →  {len(cleaned):,} chars")
    return [doc]


def load_docx(path: Path) -> list[Document]:
    """Load a Word document. Each paragraph becomes its own document."""
    try:
        import docx
    except ImportError:
        raise ImportError("pip install python-docx")

    d = docx.Document(str(path))
    paras = [p.text.strip() for p in d.paragraphs if p.text.strip()]
    if not paras:
        return []

    full_text = "\n\n".join(paras)
    doc = Document(
        page_content=_clean(full_text),
        metadata={"source": path.name, "source_path": str(path), "doc_type": "docx"},
    )
    logger.info(f"DOCX  {path.name}  →  {len(full_text):,} chars")
    return [doc]


def load_directory(corpus_dir: str | Path) -> list[Document]:
    """
    Walk a directory and load every supported file.
    Supports: .pdf, .html, .htm, .docx, .txt
    """
    corpus_dir = Path(corpus_dir)
    all_docs: list[Document] = []

    for p in sorted(corpus_dir.rglob("*")):
        if p.is_dir() or p.name.startswith("."):
            continue

        suffix = p.suffix.lower()
        try:
            if suffix == ".pdf":
                all_docs.extend(load_pdf(p))
            elif suffix in (".html", ".htm"):
                all_docs.extend(load_html(str(p), doc_type="html"))
            elif suffix == ".docx":
                all_docs.extend(load_docx(p))
            elif suffix == ".txt":
                text = p.read_text(encoding="utf-8", errors="replace")
                if text.strip():
                    all_docs.append(Document(
                        page_content=_clean(text),
                        metadata={"source": p.name, "doc_type": "txt"},
                    ))
        except Exception as e:
            logger.warning(f"skipped {p.name}: {e}")

    logger.info(f"corpus load complete  total_docs={len(all_docs)}")
    return all_docs
